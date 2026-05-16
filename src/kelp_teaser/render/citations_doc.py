"""Citations document renderer: structured Word table with one row per claim."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from kelp_teaser.schemas.citations import CitationTable

_HEADERS = ("#", "Slide", "Claim", "Source", "Verbatim Quote", "Confidence")


def render_citations_doc(table: CitationTable, out_path: Path) -> Path:
    doc = Document()

    title = doc.add_heading("Citation Audit", level=1)
    title.runs[0].font.size = Pt(18)

    doc.add_paragraph(
        "Every claim on the teaser deck below maps to a source document or web "
        "reference. Verbatim quotes are reproduced where available."
    )

    word_table = doc.add_table(rows=1, cols=len(_HEADERS))
    word_table.style = "Light Grid Accent 1"
    for i, header in enumerate(_HEADERS):
        cell = word_table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for i, row in enumerate(table.rows, start=1):
        cells = word_table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = str(row.slide_index + 1)
        cells[2].text = row.claim
        cells[3].text = row.source_id
        cells[4].text = row.verbatim_quote
        cells[5].text = row.confidence

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
