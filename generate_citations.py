import json
import os
from docx import Document
from docx.shared import Pt

# CONFIG
OUTPUT_FOLDER = "Final_Submissions"

def create_citation_doc(company_name):
    json_path = f"{company_name}_ANALYSIS.json"
    
    if not os.path.exists(json_path):
        print(f"❌ JSON not found for {company_name}")
        return

    with open(json_path, 'r') as f:
        data = json.load(f)

    doc = Document()
    
    # 1. Header
    header = doc.add_heading(f"Data Sources & Citations: {data['slide_1']['project_name']}", 0)
    
    # 2. Disclaimer (Professional Touch)
    doc.add_paragraph("The following sources were utilized by the Automated Deal Flow Agent to generate the Investment Teaser.")
    
    # 3. List Sources
    doc.add_heading('References:', level=1)
    
    if "citations" in data and isinstance(data['citations'], list):
        for source in data['citations']:
            p = doc.add_paragraph(source, style='List Bullet')
    else:
        doc.add_paragraph("No specific citations extracted.")

    # 4. Save
    output_filename = os.path.join(OUTPUT_FOLDER, f"{company_name}_Citations.docx")
    doc.save(output_filename)
    print(f"📄 Citation Doc saved: {output_filename}")

# --- TEST ---
if __name__ == "__main__":
    create_citation_doc("Test Company Name")