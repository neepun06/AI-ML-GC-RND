from docx import Document

from kelp_teaser.render.citations_doc import render_citations_doc
from kelp_teaser.schemas.citations import CitationRow, CitationTable


def _sample_table() -> CitationTable:
    return CitationTable(rows=[
        CitationRow(slide_index=0, claim="Revenue ₹450 Cr in FY24",
                    source_id="doc:report.pdf#p12",
                    verbatim_quote="Total revenue stood at ₹450 crore",
                    confidence="High"),
        CitationRow(slide_index=1, claim="35% revenue from exports",
                    source_id="web:tavily:https://example.com/about",
                    verbatim_quote="Exports comprise about a third of revenue",
                    confidence="Medium"),
    ])


def test_render_citations_doc_creates_file(tmp_path):
    out = tmp_path / "citations.docx"
    render_citations_doc(_sample_table(), out)
    assert out.exists()


def test_render_citations_doc_writes_a_table(tmp_path):
    out = tmp_path / "citations.docx"
    render_citations_doc(_sample_table(), out)
    doc = Document(str(out))
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    # Header row + 2 data rows
    assert len(table.rows) == 3
    # 5 columns: # / Slide / Claim / Source / Quote / Confidence  → 6 columns
    assert len(table.rows[0].cells) == 6


def test_render_citations_doc_includes_quote(tmp_path):
    out = tmp_path / "citations.docx"
    render_citations_doc(_sample_table(), out)
    doc = Document(str(out))
    body_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text for row in doc.tables[0].rows for cell in row.cells
    )
    full = body_text + "\n" + table_text
    assert "Total revenue stood at ₹450 crore" in full
